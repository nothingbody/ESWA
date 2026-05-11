from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "docs" / "ESWA_Final_Anonymized_Manuscript.md"
OUTPUT = ROOT / "docs" / "ESWA_Final_Anonymized_Manuscript.docx"


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    border_specs = {
        "top": ("single", "6", "808080"),
        "bottom": ("single", "6", "808080"),
        "left": ("nil", "0", "FFFFFF"),
        "right": ("nil", "0", "FFFFFF"),
        "insideH": ("nil", "0", "FFFFFF"),
        "insideV": ("nil", "0", "FFFFFF"),
    }
    for edge, (value, size, color) in border_specs.items():
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), value)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_border(cell, edge: str, value: str = "single", size: str = "6", color: str = "808080") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    tag = f"w:{edge}"
    element = borders.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        borders.append(element)
    element.set(qn("w:val"), value)
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def set_header_bottom_border(row) -> None:
    for cell in row.cells:
        set_cell_border(cell, "bottom")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_text_with_inline_format(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        else:
            run.text = part


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    set_table_borders(table)

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_text_with_inline_format(paragraph, value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                    run.font.size = Pt(9)
            if row_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    set_repeat_table_header(table.rows[0])
    set_header_bottom_border(table.rows[0])
    document.add_paragraph()


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip()
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if not all(set(cell) <= {"-", ":", " "} for cell in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def add_code_block(document: Document, code: list[str]) -> None:
    for line in code:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
    document.add_paragraph()


def add_image(document: Document, image_text: str) -> None:
    match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", image_text)
    if not match:
        return
    image_path = (ROOT / "docs" / match.group(1)).resolve()
    if not image_path.exists():
        paragraph = document.add_paragraph()
        add_text_with_inline_format(paragraph, f"[Missing figure: {image_path}]")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.4))
    document.add_paragraph()


def build_docx() -> None:
    document = Document()
    style_doc(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.title = "Anonymized manuscript"
    document.core_properties.subject = "Expert Systems with Applications submission"
    lines = INPUT.read_text(encoding="utf-8").splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if not stripped:
            idx += 1
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_with_inline_format(paragraph, stripped[2:])
            idx += 1
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
            idx += 1
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
            idx += 1
            continue
        if stripped.startswith("**") and stripped.endswith("**"):
            paragraph = document.add_paragraph()
            add_text_with_inline_format(paragraph, stripped)
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_markdown_table(lines, idx)
            add_table(document, rows)
            continue
        if stripped.startswith("```"):
            code: list[str] = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code.append(lines[idx].rstrip())
                idx += 1
            idx += 1
            add_code_block(document, code)
            continue
        if stripped.startswith("!["):
            add_image(document, stripped)
            idx += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_text_with_inline_format(paragraph, re.sub(r"^\d+\. ", "", stripped))
            idx += 1
            continue
        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_text_with_inline_format(paragraph, stripped[2:])
            idx += 1
            continue

        paragraph = document.add_paragraph()
        add_text_with_inline_format(paragraph, stripped)
        idx += 1

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
